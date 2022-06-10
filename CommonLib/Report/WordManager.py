import CommonLib.Common as common
from docx import Document
from docx.shared import RGBColor, Pt, Cm
import CommonLib.Performance.RunTest as rt
from pathlib import Path

result_Img_width = 16
result_Img_height = 5

title_Color = RGBColor(76, 89, 109)
title_size = Pt(24)
normal_font_Color = RGBColor(76, 89, 109)
normal_font_size = Pt(11)

result_report_path = 'Data/Report/ML_Result/'
result_image_path = 'Data/Image/ReportResult/'

timeSeries_Xlabel = 'Time'
timeSeries_Ylabel = 'Data'


class Cls_Paragraphs:
    """
    Word 텍스트 내에 변경할 문자내용 정의 Class
    파일명 : MLReport_Result_시퀀스.docx
    """
    def __init__(self, _text, _replaceText, _color, _fontSize, bold=False):
        self.text = _text
        self.replaceText = _replaceText
        self.color = _color
        self.fontSize = _fontSize
        self.bold = bold


def Create_Analysis_Report(titleVal, tableVal, fileName):
    """
    분석결과 리포트를 만든다
    :param titleVal: 타이틀
    :param tableVal: 내용
    :return:
    """
    try:
        word_path = common.Path_Prj_Main() + result_report_path + fileName
        document = Document(word_path)
        fileName = 'MLReport_Result_' + common.Regexp_OnlyNumberbyDate() + '.docx'
        fileFullName = common.Path_Prj_Main() + result_report_path + fileName
        for item in titleVal:
            for txtPos in document.paragraphs:
                if item.text in txtPos.text:
                    common.Word_Make_Sentense(txtPos, item)
                    break

        for item in tableVal:
            for table in document.tables:
                for col in table.columns:
                    for cell in col.cells:
                        for tablePos in cell.paragraphs:
                            if item.text in tablePos.text:
                                common.Word_Make_Sentense(tablePos, item)

        document.save(fileFullName)
        return fileFullName
    except Exception as err:
        common.exception_print(err)


def CreateReport_TimeSeries_Prophet(df, resultImagePath):
    """
    데이터 프레임에 분석결과를 추출
    :param resultImagePath:
    :param df:
    :return:
    """

    Rsquared = df.at[0, 'sts_result']
    MAE = df.at[1, 'sts_result']
    MSE = df.at[2, 'sts_result']
    MSLE = df.at[4, 'sts_result']

    titleVal = [ Cls_Paragraphs("{타이틀}", "시계열 분석 결과 보고서", title_Color, title_size, bold=True)]
    tableVal = [
        Cls_Paragraphs("{모델}", "Prophet", title_Color, normal_font_size),
        Cls_Paragraphs("{분석방법}", "daily_seasonality", title_Color, normal_font_size),
        Cls_Paragraphs("{라이브러리}", "fbprophet 패키지", title_Color, normal_font_size),
        Cls_Paragraphs("{일자}", common.ToDay(onlyDate=False), title_Color, normal_font_size),
        Cls_Paragraphs("{모델설명}", "과거에서부터 현재까지의 데이터를 바탕으로 미래에 대한 추세를 분석", title_Color, normal_font_size),
        Cls_Paragraphs("{지표1}", "R-Squared", title_Color, normal_font_size),
        Cls_Paragraphs("{지표2}", "MAE", title_Color, normal_font_size),
        Cls_Paragraphs("{지표3}", "MSE", title_Color, normal_font_size),
        Cls_Paragraphs("{지표4}", "MSLE", title_Color, normal_font_size),
        Cls_Paragraphs("{지표설명1}", Rsquared, title_Color, normal_font_size),
        Cls_Paragraphs("{지표설명2}", MAE, title_Color, normal_font_size),
        Cls_Paragraphs("{지표설명3}", MSE, title_Color, normal_font_size),
        Cls_Paragraphs("{지표설명4}", MSLE, title_Color, normal_font_size)]
    TimeSeries_Pred(tableVal, df)
    TimeSeries_Result(tableVal,df)
    report_file_name = Create_Analysis_Report(titleVal, tableVal, "MLReport_Prophet.docx")
    Result_Image_Add(report_file_name, resultImagePath)
    return report_file_name


def TimeSeries_Pred(arr, df):
    for idx in range(1, 14):
        name = df.loc[idx-1, 'type']
        result = df.loc[idx-1, 'sts_result']
        arr.append(Cls_Paragraphs("{지표T" + str(idx) + "}", str(name), title_Color, normal_font_size))
        arr.append(Cls_Paragraphs("{지표결과" + str(idx) + "}", str(result), title_Color, normal_font_size))

def TimeSeries_Result(arr, df):
    for idx in range(1, 32):
        date = df[['predic_date','predic_yhat']].tail(31).reset_index(drop=True).loc[idx - 1, 'predic_date']
        result = df[['predic_date','predic_yhat']].tail(31).reset_index(drop=True).loc[idx - 1, 'predic_yhat']
        arr.append(Cls_Paragraphs("{결과일자" + str(idx) + "}", str(date), title_Color, normal_font_size))
        arr.append(Cls_Paragraphs("{결과값" + str(idx) + "}", str(result), title_Color, normal_font_size))

# 과장님 자료
def Prophet_ParagraphsReturn(listArr, index, val, df):
    df.to_excel('Prophet_ParagraphsReturn.xlsx')
    listArr.append(Cls_Paragraphs('{지표T' + str(index) + '}', val, title_Color, normal_font_size))
    listArr.append(
        Cls_Paragraphs('{지표결과' + str(index) + '}', df.at[index - 1, 'sts_result'], title_Color, normal_font_size))




def Result_Image_Add(targetDocumentName, resultfileName):
    """
    분석결과 이미지 추가
    :param targetDocumentName: 수정 타겟 워드 파일
    :param resultfileName: 문서에 추가 할 결과이미지 파일
    :return:
    """
    try:
        if resultfileName == "":
            return

        # 결과 이미지 추가 파일
        doc = Document(targetDocumentName)
        tables = doc.tables

        imagePos = tables[2].rows[0].cells[1].add_paragraph()
        imagerun = imagePos.add_run()

        # 결과 이미지 추가
        imagerun.add_picture(resultfileName, width=Cm(result_Img_width), height=Cm(result_Img_height))
        doc.save(targetDocumentName)
    except Exception as err:
        common.exception_print(err)




def CreateReport_LinearRegression(df, resultImagePath):
    """
    데이터 프레임에 분석결과를 추출
    :param resultImagePath:
    :param df:
    :return:
    """

    Rsquared = df.at[0, 'P_MODEL_RESULT']
    MAE = df.at[1, 'P_MODEL_RESULT']
    MSE = df.at[2, 'P_MODEL_RESULT']
    MSLE = df.at[4, 'P_MODEL_RESULT']

    titleVal = [ Cls_Paragraphs("{타이틀}", "선형회귀 분석 결과 보고서", title_Color, title_size, bold=True)]
    tableVal = [
        Cls_Paragraphs("{모델}", "Ordinary Least Squares Regression", title_Color, normal_font_size),
        Cls_Paragraphs("{분석방법}", "최소자승법", title_Color, normal_font_size),
        Cls_Paragraphs("{라이브러리}", "statsmodels 패키지", title_Color, normal_font_size),
        Cls_Paragraphs("{일자}", common.ToDay(onlyDate=False), title_Color, normal_font_size),
        Cls_Paragraphs("{모델설명}", "오차를 최소화하여 회귀계수를 추정", title_Color, normal_font_size),
        Cls_Paragraphs("{지표1}", "R-Squared", title_Color, normal_font_size),
        Cls_Paragraphs("{지표2}", "MAE", title_Color, normal_font_size),
        Cls_Paragraphs("{지표3}", "MSE", title_Color, normal_font_size),
        Cls_Paragraphs("{지표4}", "MSLE", title_Color, normal_font_size),
        Cls_Paragraphs("{지표설명1}", Rsquared, title_Color, normal_font_size),
        Cls_Paragraphs("{지표설명2}", MAE, title_Color, normal_font_size),
        Cls_Paragraphs("{지표설명3}", MSE, title_Color, normal_font_size),
        Cls_Paragraphs("{지표설명4}", MSLE, title_Color, normal_font_size),
        Cls_Paragraphs("{수치1}",rt.front_Rexp(df.SUMMARY[0], 'No. Observations:', 'AIC:'), title_Color, normal_font_size),
        Cls_Paragraphs("{수치2}",rt.front_Rexp(df.SUMMARY[0], 'Df Residuals:', 'BIC:'), title_Color, normal_font_size),
        Cls_Paragraphs("{수치3}",rt.front_Rexp(df.SUMMARY[0], 'Df Model:', 'Covariance Type:'), title_Color, normal_font_size),
        Cls_Paragraphs("{수치4}",rt.front_Rexp(df.SUMMARY[0], 'R-squared:', 'Model:'), title_Color, normal_font_size),
        Cls_Paragraphs("{수치5}",rt.front_Rexp(df.SUMMARY[0], 'Adj. R-squared:', 'Method:'), title_Color, normal_font_size),
        Cls_Paragraphs("{수치6}",rt.front_Rexp(df.SUMMARY[0], 'F-statistic:', 'Date:'), title_Color, normal_font_size),
        Cls_Paragraphs("{수치7}",rt.front_Rexp(df.SUMMARY[0], 'Prob (F-statistic):', 'Time:'), title_Color, normal_font_size),
        Cls_Paragraphs("{수치8}",rt.front_Rexp(df.SUMMARY[0], 'AIC:', 'Df Residuals:'), title_Color, normal_font_size),
        Cls_Paragraphs("{수치9}",rt.front_Rexp(df.SUMMARY[0], 'BIC:', 'Df Model:'), title_Color, normal_font_size),
        Cls_Paragraphs("{수치10}",rt.front_Rexp(df.SUMMARY[0], 'Omnibus:', 'Durbin-Watson:'), title_Color, normal_font_size),
        Cls_Paragraphs("{수치11}", rt.front_Rexp(df.SUMMARY[0], 'Prob(Omnibus):', 'Jarque-Bera (JB):'), title_Color, normal_font_size),
        Cls_Paragraphs("{수치12}",rt.front_Rexp(df.SUMMARY[0], 'Skew:', 'Prob(JB):'), title_Color, normal_font_size),
        Cls_Paragraphs("{수치13}",rt.front_Rexp(df.SUMMARY[0], 'Kurtosis:', 'Cond.'), title_Color, normal_font_size),
        Cls_Paragraphs("{수치14}",rt.front_Rexp(df.SUMMARY[0], 'Durbin-Watson:', 'Prob(Omnibus):'), title_Color, normal_font_size),
        Cls_Paragraphs("{수치15}",rt.front_Rexp(df.SUMMARY[0], 'Jarque-Bera (JB): ', 'Skew:'), title_Color, normal_font_size),
        Cls_Paragraphs("{수치16}",rt.front_Rexp(df.SUMMARY[0], 'Cond. No. ', 'Notes:').replace('=','').strip(), title_Color, normal_font_size)]
    report_file_name = Create_Analysis_Report(titleVal, tableVal, "MLReport_Linear.docx")
    Result_Image_Add(report_file_name, resultImagePath)
    return report_file_name

def CreateReport_RidgeLassoRegression(df, resultImagePath):

    Rsquared = df.at[0, '평가지표']
    MAE = df.at[1, '평가지표']
    MSE = df.at[2, '평가지표']
    MSLE = df.at[4, '평가지표']

    if df.model[0] == 'ridge':
        titleVal = [Cls_Paragraphs("{타이틀}", "릿지회귀 분석 결과 보고서", title_Color, title_size, bold=True)]
        tableVal = [Cls_Paragraphs("{모델}", "Ridge", title_Color, normal_font_size)]
    else:
        titleVal = [Cls_Paragraphs("{타이틀}", "라쏘회귀 분석 결과 보고서", title_Color, title_size, bold=True)]
        tableVal = [Cls_Paragraphs("{모델}", "Lasso", title_Color, normal_font_size)]
    plus = [Cls_Paragraphs("{분석방법}", "선형회귀", title_Color, normal_font_size),
            Cls_Paragraphs("{라이브러리}", "sklearn 패키지", title_Color, normal_font_size),
            Cls_Paragraphs("{일자}", common.ToDay(onlyDate=False), title_Color, normal_font_size),
            Cls_Paragraphs("{모델설명}", "기존 선형 모델에 규제항을 추가한 회귀 모델", title_Color, normal_font_size),
            Cls_Paragraphs("{지표1}", "R-Squared", title_Color, normal_font_size),
            Cls_Paragraphs("{지표2}", "MAE", title_Color, normal_font_size),
            Cls_Paragraphs("{지표3}", "MSE", title_Color, normal_font_size),
            Cls_Paragraphs("{지표4}", "MSLE", title_Color, normal_font_size),
            Cls_Paragraphs("{지표설명1}", Rsquared, title_Color, normal_font_size),
            Cls_Paragraphs("{지표설명2}", MAE, title_Color, normal_font_size),
            Cls_Paragraphs("{지표설명3}", MSE, title_Color, normal_font_size),
            Cls_Paragraphs("{지표설명4}", MSLE, title_Color, normal_font_size)]
    tableVal.extend(plus)
    RidgeLasso_feature(tableVal, df)
    RidgeLasso_alpha(tableVal, df)
    report_file_name = Create_Analysis_Report(titleVal, tableVal, "MLReport_RidgeLasso.docx")
    Result_Image_Add(report_file_name, resultImagePath)
    return report_file_name

def RidgeLasso_feature(arr, df):
    for idx in range(1, 11):
        name = df.loc[idx-1, 'features']
        result = df.loc[idx-1, 'importances']
        arr.append(Cls_Paragraphs("{피처" + str(idx) + "}", str(name), title_Color, normal_font_size))
        arr.append(Cls_Paragraphs("{중요도" + str(idx) + "}", str(result), title_Color, normal_font_size))

def RidgeLasso_alpha(arr, df):
    for idx in range(1, 7):
        train = df.loc[idx-1, 'train_score_alpha']
        test = df.loc[idx-1, 'test_score_alpha']
        arr.append(Cls_Paragraphs("{train" + str(idx) + "}", str(train), title_Color, normal_font_size))
        arr.append(Cls_Paragraphs("{test" + str(idx) + "}", str(test), title_Color, normal_font_size))



def CreateReport_LogisticRegression(df, resultImagePath):

    Rsquared = df.at[0, '평가지표']
    MAE = df.at[1, '평가지표']
    MSE = df.at[2, '평가지표']
    MSLE = df.at[4, '평가지표']

    titleVal = [ Cls_Paragraphs("{타이틀}", "로지스틱회귀 분석 결과 보고서", title_Color, title_size, bold=True)]
    tableVal = [
        Cls_Paragraphs("{모델}", "Logistic", title_Color, normal_font_size),
        Cls_Paragraphs("{분석방법}", "선형회귀", title_Color, normal_font_size),
        Cls_Paragraphs("{라이브러리}", "sklearn 패키지", title_Color, normal_font_size),
        Cls_Paragraphs("{일자}", common.ToDay(onlyDate=False), title_Color, normal_font_size),
        Cls_Paragraphs("{모델설명}", "기존 선형 모델에 규제항을 추가한 회귀 모델", title_Color, normal_font_size),
        Cls_Paragraphs("{지표1}", "R-Squared", title_Color, normal_font_size),
        Cls_Paragraphs("{지표2}", "MAE", title_Color, normal_font_size),
        Cls_Paragraphs("{지표3}", "MSE", title_Color, normal_font_size),
        Cls_Paragraphs("{지표4}", "MSLE", title_Color, normal_font_size),
        Cls_Paragraphs("{지표설명1}", Rsquared, title_Color, normal_font_size),
        Cls_Paragraphs("{지표설명2}", MAE, title_Color, normal_font_size),
        Cls_Paragraphs("{지표설명3}", MSE, title_Color, normal_font_size),
        Cls_Paragraphs("{지표설명4}", MSLE, title_Color, normal_font_size),
        Cls_Paragraphs("{수치1}", rt.front_Rexp(df.SUMMARY[0], 'Method:', 'Df Model:'), title_Color,normal_font_size),
        Cls_Paragraphs("{수치2}", rt.front_Rexp(df.SUMMARY[0], 'Covariance Type:', 'LLR p-value:'), title_Color, normal_font_size),
        Cls_Paragraphs("{수치3}", rt.front_Rexp(df.SUMMARY[0], 'No. Observations:', 'Model:'), title_Color,normal_font_size),
        Cls_Paragraphs("{수치4}", rt.front_Rexp(df.SUMMARY[0], 'Df Residuals:', 'Method:'), title_Color, normal_font_size),
        Cls_Paragraphs("{수치5}", rt.front_Rexp(df.SUMMARY[0], 'Df Model:', 'Date:'), title_Color,normal_font_size),
        Cls_Paragraphs("{수치6}", rt.front_Rexp(df.SUMMARY[0], 'Pseudo R-squ.:', 'Time:'), title_Color, normal_font_size),
        Cls_Paragraphs("{수치7}", rt.front_Rexp(df.SUMMARY[0], 'Log-Likelihood:', 'converged:'), title_Color,normal_font_size),
        Cls_Paragraphs("{수치8}", rt.front_Rexp(df.SUMMARY[0], 'LL-Null:', 'Covariance Type:'), title_Color, normal_font_size),
        Cls_Paragraphs("{수치9}", rt.front_Rexp(df.SUMMARY[0], 'LLR p-value:', 'coef').replace('=','').strip(), title_Color, normal_font_size)]
    Logistic_odds(tableVal, df)
    report_file_name = Create_Analysis_Report(titleVal, tableVal, "MLReport_logistic.docx")
    Result_Image_Add(report_file_name, resultImagePath)
    return report_file_name

def Logistic_odds(arr, df):
    for idx in range(1, 10):
        name = df.loc[idx-1, 'Params']
        result = df.loc[idx-1, 'Odds Ratio']
        arr.append(Cls_Paragraphs("{피처" + str(idx) + "}", str(name), title_Color, normal_font_size))
        arr.append(Cls_Paragraphs("{오즈비" + str(idx) + "}", str(result), title_Color, normal_font_size))


def CreateReport_RandomForest(df, resultImagePath):

    Rsquared = df.at[0, '평가지표']
    MAE = df.at[1, '평가지표']
    MSE = df.at[2, '평가지표']
    MSLE = df.at[4, '평가지표']

    titleVal = [ Cls_Paragraphs("{타이틀}", "랜던포레스트 분석 결과 보고서", title_Color, title_size, bold=True)]
    tableVal = [
        Cls_Paragraphs("{모델}", "RandomForest", title_Color, normal_font_size),
        Cls_Paragraphs("{분석방법}", "분류", title_Color, normal_font_size),
        Cls_Paragraphs("{라이브러리}", "sklearn 패키지", title_Color, normal_font_size),
        Cls_Paragraphs("{일자}", common.ToDay(onlyDate=False), title_Color, normal_font_size),
        Cls_Paragraphs("{모델설명}", "기존 선형 모델에 규제항을 추가한 회귀 모델", title_Color, normal_font_size),
        Cls_Paragraphs("{지표1}", "R-Squared", title_Color, normal_font_size),
        Cls_Paragraphs("{지표2}", "MAE", title_Color, normal_font_size),
        Cls_Paragraphs("{지표3}", "MSE", title_Color, normal_font_size),
        Cls_Paragraphs("{지표4}", "MSLE", title_Color, normal_font_size),
        Cls_Paragraphs("{지표설명1}", Rsquared, title_Color, normal_font_size),
        Cls_Paragraphs("{지표설명2}", MAE, title_Color, normal_font_size),
        Cls_Paragraphs("{지표설명3}", MSE, title_Color, normal_font_size),
        Cls_Paragraphs("{지표설명4}", MSLE, title_Color, normal_font_size),
        Cls_Paragraphs("{정확도}", df['정확도'][0], title_Color,normal_font_size)]
    RandomForest_Importance(tableVal, df)
    report_file_name = Create_Analysis_Report(titleVal, tableVal, "MLReport_RandomForest.docx")
    Result_Image_Add(report_file_name, resultImagePath)
    return report_file_name

def RandomForest_Importance(arr, df):
    for idx in range(1, 10):
        name = df.loc[idx-1, 'feature']
        result = df.loc[idx-1, 'importance']
        arr.append(Cls_Paragraphs("{피처" + str(idx) + "}", str(name), title_Color, normal_font_size))
        arr.append(Cls_Paragraphs("{중요도" + str(idx) + "}", str(result), title_Color, normal_font_size))

# if __name__ == '__main__':
#     try:
#         Rsquared = 123
#         MAE = 325
#         MSE = 4565
#         MSLE = 1233
#
#         titleVal = [Cls_Paragraphs("{타이틀}", "선형회귀 분석 결과 보고서", title_Color, title_size, bold=True)]
#         tableVal = [
#             Cls_Paragraphs("{모델}", "Ordinary Least Squares Regression", title_Color, normal_font_size),
#             Cls_Paragraphs("{분석방법}", "최소자승법", title_Color, normal_font_size),
#             Cls_Paragraphs("{라이브러리}", "statsmodels 패키지", title_Color, normal_font_size),
#             Cls_Paragraphs("{일자}", common.ToDay(onlyDate=False), title_Color, normal_font_size),
#             Cls_Paragraphs("{모델설명}", "오차를 최소화하여 회귀계수를 추정", title_Color, normal_font_size),
#             Cls_Paragraphs("{지표1}", "R-Squared", title_Color, normal_font_size),
#             Cls_Paragraphs("{지표2}", "MAE", title_Color, normal_font_size),
#             Cls_Paragraphs("{지표3}", "MSE", title_Color, normal_font_size),
#             Cls_Paragraphs("{지표4}", "MSLE", title_Color, normal_font_size),
#             Cls_Paragraphs("{지표설명1}", Rsquared, title_Color, normal_font_size),
#             Cls_Paragraphs("{지표설명2}", MAE, title_Color, normal_font_size),
#             Cls_Paragraphs("{지표설명3}", MSE, title_Color, normal_font_size),
#             Cls_Paragraphs("{지표설명4}", MSLE, title_Color, normal_font_size),
#             Cls_Paragraphs("{분석결과설명}",
#                            """
#                            ①결과에 대한 분석전, Durbin watson을 확인한 결과 1.780으로 2에 근접하여
#                            다중회귀분석모형에 적합하다고 판단하였다.
#                            ②그리고 유의확률이 .000으로 경로 중
#                            한가지 이상이 유효할 것이라것을 확인할 수 있었다.
#                            ③ 또한 변수간 상관관계(R=.712)로 확인되었고
#                            ④ 공차와 VIF 각각 0.1이상 10미만으로 다중공선성이 없는 것으로 확인되었다.
#                            ⑤ 다음으로 각 경로에 유의성을 확인한 결과 신체자발성을 제외한
#                            사회자발성(p<.001)과 인지자발성(p<.001)이 기질에 미치는 영향이 유효한 것으로 확인되었다.
#                            ⑥ 유의한 변수에 대한 비표준화계수를 확인한 결과 사회자발셩(B=.283), 인지자발성(B=.369) 모두 양수로써
#                            사회자발성과 인지자발성이 향상될수록 기질이 높아진 다는 것을 알 수 있었다.
#                            ⑦또한, 사회자발성(베타=.375)과 인지자발성(베타=.445)이
#                            기질에 미치는 영향력은 인지자발성, 사회자발성 순임을 알 수 있었다.
#                            ⑧ 마지막으로 독립변수에 의해 종속변수가 설명되는 설명력은 50.7%임을 확인할 수 있었다.
#                            """, title_Color, Pt(9))]
#
#         word_path = 'D:/26.DataAnalisyPrj/ProxyServer/FastAPIServer/Data/Report/ML_Result/MLReport.docx'
#         document = Document(word_path)
#         fileName = 'MLReport_Result_' + common.Regexp_OnlyNumberbyDate() + '.docx'
#         fileFullName = common.Path_Prj_Main() + result_report_path + fileName
#         for item in titleVal:
#             for txtPos in document.paragraphs:
#                 if item.text in txtPos.text:
#                     # common.Word_Make_Sentense(txtPos, item.text, item.replaceText, item.color, item.fontSize, item.bold, ITEM.STRIKE)
#                     common.Word_Make_Sentense(txtPos, item)
#                     break
#
#         for item in tableVal:
#             for table in document.tables:
#                 for col in table.columns:
#                     for cell in col.cells:
#                         for tablePos in cell.paragraphs:
#                             if item.text in tablePos.text:
#                                 common.Word_Make_Sentense(tablePos, item.text, item.replaceText, item.color,
#                                                           item.fontSize, item.bold)
#
#         document.save(fileFullName)
#     except Exception as err:
#         common.exception_print(err)

# def CreateReport_RidgeRegression(df, resultImagePath):
#     """
#     데이터 프레임에 분석결과를 추출
#     :param resultImagePath:
#     :param df:
#     :return:
#     """
#
#     Rsquared = df.at[0, '평가지표']
#     MAE = df.at[1, '평가지표']
#     MSE = df.at[2, '평가지표']
#     MSLE = df.at[4, '평가지표']
#
#     titleVal = [ Cls_Paragraphs("{타이틀}", "릿지회귀 분석 결과 보고서", title_Color, title_size, bold=True)]
#     tableVal = [
#         Cls_Paragraphs("{모델}", "Ridge", title_Color, normal_font_size),
#         Cls_Paragraphs("{분석방법}", "선형회귀", title_Color, normal_font_size),
#         Cls_Paragraphs("{라이브러리}", "sklearn 패키지", title_Color, normal_font_size),
#         Cls_Paragraphs("{일자}", common.ToDay(onlyDate=False), title_Color, normal_font_size),
#         Cls_Paragraphs("{모델설명}", "기존 선형 모델에 규제항을 추가한 회귀 모델", title_Color, normal_font_size),
#         Cls_Paragraphs("{지표1}", "R-Squared", title_Color, normal_font_size),
#         Cls_Paragraphs("{지표2}", "MAE", title_Color, normal_font_size),
#         Cls_Paragraphs("{지표3}", "MSE", title_Color, normal_font_size),
#         Cls_Paragraphs("{지표4}", "MSLE", title_Color, normal_font_size),
#         Cls_Paragraphs("{지표설명1}", Rsquared, title_Color, normal_font_size),
#         Cls_Paragraphs("{지표설명2}", MAE, title_Color, normal_font_size),
#         Cls_Paragraphs("{지표설명3}", MSE, title_Color, normal_font_size),
#         Cls_Paragraphs("{지표설명4}", MSLE, title_Color, normal_font_size)]
#     Ridge_feature(tableVal, df)
#     Ridge_alpha(tableVal, df)
#     report_file_name = Create_Analysis_Report(titleVal, tableVal, "MLReport_Ridge.docx")
#     Result_Image_Add(report_file_name, resultImagePath)
#     return report_file_name
#
# def Ridge_feature(arr, df):
#     for idx in range(1, 12):
#         name = df.loc[idx-1, 'features']
#         result = df.loc[idx-1, 'importances']
#         arr.append(Cls_Paragraphs("{피처" + str(idx) + "}", str(name), title_Color, normal_font_size))
#         arr.append(Cls_Paragraphs("{중요도" + str(idx) + "}", str(result), title_Color, normal_font_size))
#
# def Ridge_alpha(arr, df):
#     for idx in range(1, 7):
#         train = df.loc[idx-1, 'train_score_alpha']
#         test = df.loc[idx-1, 'test_score_alpha']
#         arr.append(Cls_Paragraphs("{train" + str(idx) + "}", str(train), title_Color, normal_font_size))
#         arr.append(Cls_Paragraphs("{test" + str(idx) + "}", str(test), title_Color, normal_font_size))